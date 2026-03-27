"""
app/reports/protocol_docx.py

Generate "ΠΡΩΤΟΚΟΛΛΟ" as DOCX bytes using a Word template and placeholder
replacement.

SOURCE OF TRUTH
---------------
This implementation is aligned to:
- current uploaded `protocol_template.docx`
- current Procurement / ServiceUnit / Committee / Directory related models
- the canonical procurement-type rule:
  if any procurement material line has `is_service == True`
  -> services wording
  otherwise -> goods/materials

IMPORTANT DIRECTORY RULE
------------------------
For the materials-only closing block, `{{HANDLER_NAME}}` must resolve to the
director of the selected handler directory, not the department head and not the
handler person.

Resolution order:
1. procurement.handler_assignment.directory.director
2. procurement.handler_assignment.directory.director_personnel
3. procurement.handler_assignment.directory.director_personnel_id -> lookup
4. em dash fallback

CURRENT YEAR RULE
-----------------
`{{CURRENT_YEAR}}` resolves to the 2-digit year:
1. procurement.materials_receipt_date.year
2. current system year

TEMPLATE HANDLING RULE
----------------------
The protocol template is the layout source of truth. This generator must:
- replace placeholders only
- preserve the template's alignment / tabs / table layout
- avoid rebuilding the committee-signature structure
"""

from __future__ import annotations

from contextlib import nullcontext
from dataclasses import dataclass
from datetime import date
from decimal import Decimal
from io import BytesIO
from pathlib import Path
from typing import Any, Iterable, Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

from ..models.organization import Personnel
from .common.amounts import money_plain
from .common.docx_utils import (
    clear_table_body_keep_header,
    set_cell_alignment,
    set_global_font_arial_12,
)
from .common.domain import (
    resolve_is_services,
    resolve_service_unit_place,
    winner_supplier_line,
)
from .common.formatting import format_date_ddmmyyyy, safe_text, upper_service_name
from .instrumentation import ReportInstrumentation


def _timed(
    instrumentation: Optional[ReportInstrumentation],
    detail_name: str,
    **extra: Any,
):
    if instrumentation is None:
        return nullcontext()
    return instrumentation.timed_detail(detail_name, **extra)


def _template_path() -> Path:
    """
    Resolve the DOCX template path for the protocol document.
    """
    app_dir = Path(__file__).resolve().parents[1]
    return app_dir / "templates" / "docx" / "protocol_template.docx"


def _sorted_materials(procurement: Any) -> list[Any]:
    """
    Preserve source ordering by default.
    """
    return list(getattr(procurement, "materials", []) or [])


def _person_full_name(person: Any) -> str:
    """
    Resolve a person display name conservatively.
    """
    if person is None:
        return "—"

    full_name_method = getattr(person, "full_name", None)
    if callable(full_name_method):
        try:
            value = safe_text(full_name_method())
            if value:
                return value
        except Exception:
            pass

    rank = safe_text(getattr(person, "rank", None), default="")
    last_name = safe_text(getattr(person, "last_name", None), default="")
    first_name = safe_text(getattr(person, "first_name", None), default="")
    specialty = safe_text(getattr(person, "specialty", None), default="")

    combined = " ".join(
        part for part in [rank, specialty, first_name, last_name] if part
    ).strip()
    if combined:
        return combined

    name = safe_text(getattr(person, "name", None), default="")
    if name:
        return name

    return "—"


def _resolve_directory_director_name(procurement: Any) -> str:
    """
    Resolve the selected handler directory director.
    """
    assignment = getattr(procurement, "handler_assignment", None)
    if assignment is not None:
        directory = getattr(assignment, "directory", None)
        if directory is not None:
            director = getattr(directory, "director", None)
            if director is not None:
                return _person_full_name(director)

            director_personnel = getattr(directory, "director_personnel", None)
            if director_personnel is not None:
                return _person_full_name(director_personnel)

            director_personnel_id = getattr(directory, "director_personnel_id", None)
            if director_personnel_id:
                person = Personnel.query.get(director_personnel_id)
                if person is not None:
                    return _person_full_name(person)

    return "—"


def _committee_member(procurement: Any, attr_name: str) -> str:
    committee = getattr(procurement, "committee", None)
    if committee is None:
        return "—"
    return _person_full_name(getattr(committee, attr_name, None))


def _committee_identity_text(procurement: Any) -> str:
    committee = getattr(procurement, "committee", None)
    if committee is None:
        return "—"
    return safe_text(getattr(committee, "identity_text", None), default="—")


def _protocol_title(is_services: bool) -> str:
    return (
        "ΠΡΩΤΟΚΟΛΛΟ ΠΟΣΟΤΙΚΗΣ ΚΑΙ ΠΟΙΟΤΙΚΗΣ ΠΑΡΟΧΗΣ ΥΠΗΡΕΣΙΩΝ"
        if is_services
        else "ΠΡΩΤΟΚΟΛΛΟ ΠΟΣΟΤΙΚΗΣ ΚΑΙ ΠΟΙΟΤΙΚΗΣ ΠΡΟΜΗΘΕΙΑΣ ΥΛΙΚΩΝ"
    )


def _intro_items_label(is_services: bool) -> str:
    return "υπηρεσιών" if is_services else "υλικών"


def _delivery_verb(is_services: bool) -> str:
    return "παρασχέθηκαν" if is_services else "παραλήφθηκαν"


def _analysis_title(is_services: bool) -> str:
    return (
        "ΑΝΑΛΥΣΗ ΠΑΡΟΧΗΣ ΥΠΗΡΕΣΙΩΝ"
        if is_services
        else "ΑΝΑΛΥΣΗ ΠΡΟΜΗΘΕΙΑΣ ΥΛΙΚΩΝ"
    )


def _decision_text(is_services: bool) -> str:
    if is_services:
        return (
            "Η επιτροπή, αφού έλαβε υπόψη τα ανωτέρω, αποφαίνεται παμψηφεί ότι "
            "οι υπηρεσίες που παρασχέθηκαν, πληρούν τους όρους της σύμβασης και "
            "εισηγείται την οριστική παραλαβή αυτών."
        )

    return (
        "Η επιτροπή, αφού έλαβε υπόψη τα ανωτέρω, αποφαίνεται παμψηφεί ότι "
        "τα υλικά που παραλήφθηκαν, πληρούν τους όρους της σύμβασης και "
        "εισηγείται την οριστική παραλαβή αυτών."
    )


def _current_year_text(procurement: Any) -> str:
    materials_receipt_date = getattr(procurement, "materials_receipt_date", None)
    if materials_receipt_date is not None and getattr(materials_receipt_date, "year", None):
        return str(materials_receipt_date.year)
    return str(date.today().year)


def _current_year_two_digits(procurement: Any) -> str:
    return _current_year_text(procurement)[-2:]


def _find_protocol_table(doc: Document):
    """
    Find the main protocol analysis table.

    Expected shape in current template:
    - 6 columns
    - header includes:
      Α/Α | ΠΕΡΙΓΡΑΦΗ | NSN | Μ/Μ | ΣΥΝ. ΠΟΣ. | ΤΙΜΗ ΜΟΝ. (€)
    """
    for table in doc.tables:
        if len(table.columns) != 6 or not table.rows:
            continue

        header = " ".join(cell.text.strip() for cell in table.rows[0].cells).upper()
        if (
            "Α/Α" in header
            and "ΠΕΡΙΓΡΑΦΗ" in header
            and "NSN" in header
            and "Μ/Μ" in header
            and "ΣΥΝ. ΠΟΣ." in header
            and "ΤΙΜΗ ΜΟΝ." in header
        ):
            return table

    return None


def _fill_protocol_table(table, procurement: Any) -> None:
    """
    Fill the protocol analysis table with one row per procurement material line.
    """
    clear_table_body_keep_header(table, header_rows=1)

    lines = _sorted_materials(procurement)

    if not lines:
        row = table.add_row().cells
        row[0].text = "—"
        row[1].text = "Δεν υπάρχουν γραμμές υλικών/υπηρεσιών."
        row[2].text = "—"
        row[3].text = "—"
        row[4].text = "—"
        row[5].text = "—"
        for cell in row:
            set_cell_alignment(
                cell,
                horizontal=WD_ALIGN_PARAGRAPH.CENTER,
                vertical=WD_ALIGN_VERTICAL.CENTER,
            )
        return

    for idx, line in enumerate(lines, start=1):
        row = table.add_row().cells
        row[0].text = str(idx)
        row[1].text = safe_text(getattr(line, "description", None))
        row[2].text = safe_text(getattr(line, "nsn", None))
        row[3].text = safe_text(getattr(line, "unit", None))

        qty_value = getattr(line, "quantity", None)
        row[4].text = money_plain(qty_value) if qty_value is not None else "—"
        row[5].text = money_plain(getattr(line, "unit_price", None))

        for cell in row:
            set_cell_alignment(
                cell,
                horizontal=WD_ALIGN_PARAGRAPH.CENTER,
                vertical=WD_ALIGN_VERTICAL.CENTER,
            )


def _iter_body_blocks(doc: Document):
    """
    Iterate document body elements preserving body order.
    """
    body = doc.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield ("paragraph", child)
        elif isinstance(child, CT_Tbl):
            yield ("table", child)


def _remove_xml_element(element) -> None:
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _iter_all_paragraphs(doc: Document) -> Iterable[Paragraph]:
    for paragraph in doc.paragraphs:
        yield paragraph

    for table in doc.tables:
        yield from _iter_table_paragraphs(table)

    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            yield paragraph
        for table in section.header.tables:
            yield from _iter_table_paragraphs(table)

        for paragraph in section.footer.paragraphs:
            yield paragraph
        for table in section.footer.tables:
            yield from _iter_table_paragraphs(table)


def _iter_table_paragraphs(table: Table) -> Iterable[Paragraph]:
    for row in table.rows:
        for cell in row.cells:
            yield from _iter_cell_paragraphs(cell)


def _iter_cell_paragraphs(cell: _Cell) -> Iterable[Paragraph]:
    for paragraph in cell.paragraphs:
        yield paragraph
    for table in cell.tables:
        yield from _iter_table_paragraphs(table)


def _replace_text_in_paragraph(paragraph: Paragraph, old: str, new: str) -> None:
    if old not in paragraph.text:
        return

    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)

    if old in paragraph.text:
        paragraph.text = paragraph.text.replace(old, new)


def _replace_placeholders_everywhere_local(doc: Document, mapping: dict[str, str]) -> None:
    """
    Replace placeholders across body, tables, headers, and footers.
    """
    for paragraph in _iter_all_paragraphs(doc):
        if not paragraph.text:
            continue
        for placeholder, value in mapping.items():
            if placeholder in paragraph.text:
                _replace_text_in_paragraph(paragraph, placeholder, value)


def _fix_known_template_typos(doc: Document) -> None:
    """
    Patch known current-template typos before placeholder replacement.

    Important:
    - only replace the exact wrong placeholder token
    - do not rewrite whole paragraph layout
    """
    wrong = "{{COMMITTEE_MEMBER1}}"
    right = "{{COMMITTEE_MEMBER3}}"

    for paragraph in _iter_all_paragraphs(doc):
        text = paragraph.text or ""
        if text.strip().startswith("γ.") and wrong in text:
            _replace_text_in_paragraph(paragraph, wrong, right)


def _strip_marker_paragraphs_and_toggle_materials_block(
    doc: Document,
    *,
    keep_materials_block: bool,
) -> None:
    """
    Remove materials markers and optionally remove the whole materials block.

    Rules:
    - when materials block is kept: remove only the markers
    - when materials block is not kept: remove everything between markers
      from the main document body, including tables
    """
    start_marker = "{{#IF_MATERIALS}}"
    end_marker = "{{/IF_MATERIALS}}"

    inside_block = False

    for kind, element in list(_iter_body_blocks(doc)):
        if kind == "paragraph":
            text = "".join(element.itertext())

            has_start = start_marker in text
            has_end = end_marker in text

            if has_start and not keep_materials_block:
                inside_block = True
                _remove_xml_element(element)
                continue

            if has_end and not keep_materials_block:
                inside_block = False
                _remove_xml_element(element)
                continue

            if inside_block and not keep_materials_block:
                _remove_xml_element(element)
                continue

            if has_start or has_end:
                paragraph = Paragraph(element, doc)
                if has_start:
                    _replace_text_in_paragraph(paragraph, start_marker, "")
                if has_end:
                    _replace_text_in_paragraph(paragraph, end_marker, "")

        elif kind == "table" and inside_block and not keep_materials_block:
            _remove_xml_element(element)

    _remove_marker_tokens_everywhere(doc)


def _remove_marker_tokens_everywhere(doc: Document) -> None:
    """
    Final safety pass to ensure IF markers never remain visible anywhere.
    """
    for paragraph in _iter_all_paragraphs(doc):
        if not paragraph.text:
            continue
        if "{{#IF_MATERIALS}}" in paragraph.text:
            _replace_text_in_paragraph(paragraph, "{{#IF_MATERIALS}}", "")
        if "{{/IF_MATERIALS}}" in paragraph.text:
            _replace_text_in_paragraph(paragraph, "{{/IF_MATERIALS}}", "")


def _resolve_total_amount(procurement: Any, analysis: Optional[dict[str, Any]] = None) -> Decimal:
    """
    Resolve the filename amount.

    Business rule:
    - use General Total, not Final Payable Total
    """
    analysis = analysis or {}

    for key in ("sum_total", "total", "grand_total"):
        value = analysis.get(key)
        if value not in (None, ""):
            try:
                return Decimal(str(value))
            except Exception:
                pass

    for attr_name in ("sum_total", "total", "grand_total", "total_amount"):
        value = getattr(procurement, attr_name, None)
        if value not in (None, ""):
            try:
                return Decimal(str(value))
            except Exception:
                pass

    running = Decimal("0")
    found = False
    for line in _sorted_materials(procurement):
        for attr_name in ("total_with_vat", "total_post_vat", "total_pre_vat", "total"):
            value = getattr(line, attr_name, None)
            if value not in (None, ""):
                try:
                    running += Decimal(str(value))
                    found = True
                    break
                except Exception:
                    pass

    if found:
        return running

    return Decimal("0")


def _format_filename_amount(amount: Decimal) -> str:
    q = amount.quantize(Decimal("0.01"))
    s = f"{q:.2f}".replace(".", ",")
    return f"{s}Ε"


@dataclass(frozen=True)
class ProtocolConstants:
    """
    Future-proof constants container for protocol generation.
    """
    pass


def build_protocol_docx(
    *,
    procurement: Any,
    service_unit: Any,
    winner: Optional[Any],
    analysis: dict[str, Any],
    constants: Optional[ProtocolConstants] = None,
    instrumentation: Optional[ReportInstrumentation] = None,
) -> bytes:
    """
    Build the protocol DOCX and return it as bytes.
    """
    _ = constants

    template_path = _template_path()
    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")

    is_services = resolve_is_services(procurement)

    with _timed(instrumentation, "load_template"):
        doc = Document(str(template_path))

    with _timed(instrumentation, "fix_template_typos"):
        _fix_known_template_typos(doc)

    with _timed(instrumentation, "build_mapping"):
        mapping: dict[str, str] = {
            "{{PROCUREMENT_PROTOCOL_NUMBER}}": safe_text(
                getattr(procurement, "protocol_number", None)
            ),
            "{{CURRENT_YEAR}}": _current_year_two_digits(procurement),
            "{{SERVICE_UNIT_NAME}}": upper_service_name(
                safe_text(getattr(service_unit, "description", None), default="—")
            ),
            "{{PROTOCOL_TITLE}}": _protocol_title(is_services),
            "{{PROCUREMENT_SERVICE_UNIT_PLACE}}": resolve_service_unit_place(service_unit),
            "{{PROCUREMENT_MATERIALS_RECEIPT_DATE}}": format_date_ddmmyyyy(
                getattr(procurement, "materials_receipt_date", None)
            ),
            "{{COMMITTEE_MEMBER1}}": _committee_member(procurement, "president"),
            "{{COMMITTEE_MEMBER2}}": _committee_member(procurement, "member1"),
            "{{COMMITTEE_MEMBER3}}": _committee_member(procurement, "member2"),
            "{{PROCUREMENT_COMMITTEE_IDENTITY_TEXT}}": _committee_identity_text(procurement),
            "{{PROTOCOL_INTRO_ITEMS_LABEL}}": _intro_items_label(is_services),
            "{{PROTOCOL_DELIVERY_VERB}}": _delivery_verb(is_services),
            "{{WINNER_SUPPLIER_LINE}}": winner_supplier_line(winner),
            "{{PROCUREMENT_HOP_APPROVAL}}": safe_text(
                getattr(procurement, "hop_approval", None)
            ),
            "{{PROTOCOL_ANALYSIS_TITLE}}": _analysis_title(is_services),
            "{{PROTOCOL_DECISION_TEXT}}": _decision_text(is_services),
            "{{service.commander}}": safe_text(
                getattr(service_unit, "commander", None),
                default="—",
            ),
            "{{COMMANDER_ROLE_TYPE}}": safe_text(
                getattr(service_unit, "commander_role_type", None),
                default="—",
            ),
            "{{SERVICE_UNIT_SUPPLY_OFFICER}}": safe_text(
                getattr(service_unit, "supply_officer", None),
                default="—",
            ),
            "{{HANDLER_NAME}}": _resolve_directory_director_name(procurement),
            "{{ML_NO}}": "",
            "{{ML_DESC}}": "",
            "{{ML_NSN}}": "",
            "{{ML_UNIT}}": "",
            "{{ML_QTY}}": "",
            "{{ML_UNIT_PRICE}}": "",
        }

    with _timed(
        instrumentation,
        "toggle_materials_block",
        keep_materials_block=not is_services,
    ):
        _strip_marker_paragraphs_and_toggle_materials_block(
            doc,
            keep_materials_block=not is_services,
        )

    with _timed(
        instrumentation,
        "replace_placeholders_body_headers_footers",
        placeholders=len(mapping),
    ):
        _replace_placeholders_everywhere_local(doc, mapping)

    with _timed(instrumentation, "locate_tables"):
        protocol_table = _find_protocol_table(doc)

    if protocol_table is not None:
        with _timed(
            instrumentation,
            "fill_protocol_table",
            materials_count=len(_sorted_materials(procurement)),
        ):
            _fill_protocol_table(protocol_table, procurement)

    with _timed(instrumentation, "set_global_font"):
        set_global_font_arial_12(doc)

    with _timed(instrumentation, "save_docx"):
        buffer = BytesIO()
        doc.save(buffer)

    return buffer.getvalue()


def build_protocol_filename(
    *,
    procurement: Any,
    winner: Optional[Any],
    is_services: bool,
    analysis: Optional[dict[str, Any]] = None,
) -> str:
    """
    Build a human-readable filename for the generated protocol document.

    Expected pattern:
    Πρωτόκολλο Ποιοτικής και Ποσοτικής Παραλαβής <kind> <protocol>_<yy> <supplier> <amount>Ε.docx

    Amount rule:
    - use General Total, not Final Payable Total
    """
    kind = "Παροχής Υπηρεσιών" if is_services else "Προμήθειας Υλικών"
    supplier_name = safe_text(getattr(winner, "name", None), default="—")
    protocol_number = safe_text(getattr(procurement, "protocol_number", None), default="—")
    protocol_with_year = f"{protocol_number}_{_current_year_two_digits(procurement)}"
    amount_text = _format_filename_amount(_resolve_total_amount(procurement, analysis))

    return (
        f"Πρωτόκολλο Ποιοτικής και Ποσοτικής Παραλαβής "
        f"{kind} {protocol_with_year} {supplier_name} {amount_text}.docx"
    )