"""
app/reports/invitation_docx.py

Generate "ΠΡΟΣΚΛΗΣΗ" as DOCX bytes using a Word template and placeholder
replacement.

SOURCE OF TRUTH
---------------
This implementation is aligned to:
- current uploaded `invitation_template.docx`
- current Procurement / ServiceUnit / Supplier related models

IMPORTANT TABLE RULE
--------------------
The invitation document must render the materials/services table with:
- one document row per procurement line
- not multiline text packed into a single row

BUSINESS RULE
-------------
The document switches wording dynamically between:
- παροχή υπηρεσιών
- προμήθεια υλικών

Canonical project rule:
- if any procurement material line has `is_service == True`
  -> services wording
- otherwise
  -> goods/materials wording

KNOWN MODEL LIMITATIONS
-----------------------
The current source-of-truth models do NOT provide dedicated fields for:
- explicit service-unit place for signing/sending
- richer recipient metadata beyond supplier fields already present

Therefore:
- {{SERVICE_UNIT_PLACE}} resolves conservatively to "—"
"""

from __future__ import annotations

import unicodedata
from dataclasses import dataclass
from datetime import date, datetime
from io import BytesIO
from pathlib import Path
from typing import Any, Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


# ---------------------------------------------------------------------------
# Generic helpers
# ---------------------------------------------------------------------------

def _safe(value: Any, default: str = "—") -> str:
    text = ("" if value is None else str(value)).strip()
    return text if text else default


def _to_decimal(value: Any):
    from decimal import Decimal

    try:
        return Decimal(str(value or "0"))
    except Exception:
        return Decimal("0.00")


def _money_plain(value: Any) -> str:
    amount = _to_decimal(value).quantize(_to_decimal("0.01"))
    text = f"{amount:,.2f}"
    return text.replace(",", "X").replace(".", ",").replace("X", ".")


def _upper_no_accents(value: Any, default: str = "—") -> str:
    text = _safe(value, default=default)
    normalized = unicodedata.normalize("NFD", text)
    no_marks = "".join(ch for ch in normalized if not unicodedata.combining(ch))
    return unicodedata.normalize("NFC", no_marks).upper()


def _lower_preserve_accents(value: Any, default: str = "—") -> str:
    """
    Return lowercase text while preserving accents/diacritics.
    """
    text = _safe(value, default=default)
    return text.lower()


def _format_date(value: Any, default: str = "—") -> str:
    if value is None:
        return default

    if isinstance(value, date):
        return value.strftime("%d/%m/%Y")

    try:
        return value.strftime("%d/%m/%Y")
    except Exception:
        text = str(value).strip()
        return text if text else default


def _short_date_el(value: Any | None = None) -> str:
    months = {
        1: "Ιαν",
        2: "Φεβ",
        3: "Μαρ",
        4: "Απρ",
        5: "Μαϊ",
        6: "Ιουν",
        7: "Ιουλ",
        8: "Αυγ",
        9: "Σεπ",
        10: "Οκτ",
        11: "Νοε",
        12: "Δεκ",
    }

    dt = value or datetime.now()

    try:
        day = int(dt.day)
        month = int(dt.month)
        year_2d = int(dt.year) % 100
    except Exception:
        dt = datetime.now()
        day = dt.day
        month = dt.month
        year_2d = dt.year % 100

    month_label = months.get(month, "")
    return f"{day:02d} {month_label} {year_2d:02d}".strip()


# ---------------------------------------------------------------------------
# Template helpers
# ---------------------------------------------------------------------------

def _template_path() -> Path:
    app_dir = Path(__file__).resolve().parents[1]
    return app_dir / "templates" / "docx" / "invitation_template.docx"


def _set_global_font_arial_12(doc: Document) -> None:
    try:
        style = doc.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(12)
    except Exception:
        pass


def _paragraph_runs_text(paragraph) -> tuple[str, list[tuple[int, int, int]]]:
    pieces: list[str] = []
    positions: list[tuple[int, int, int]] = []

    cursor = 0
    for idx, run in enumerate(paragraph.runs):
        text = run.text or ""
        pieces.append(text)
        start = cursor
        end = start + len(text)
        positions.append((idx, start, end))
        cursor = end

    return "".join(pieces), positions


def _find_run_index_at_offset(positions: list[tuple[int, int, int]], offset: int) -> int:
    for run_index, start, end in positions:
        if start <= offset < end:
            return run_index

    if positions and offset == positions[-1][2]:
        return positions[-1][0]

    return -1


def _replace_placeholder_once_in_paragraph(paragraph, placeholder: str, replacement: str) -> bool:
    full_text, positions = _paragraph_runs_text(paragraph)
    if not full_text or placeholder not in full_text:
        return False

    start = full_text.find(placeholder)
    end = start + len(placeholder)

    start_run_idx = _find_run_index_at_offset(positions, start)
    end_run_idx = _find_run_index_at_offset(positions, end - 1)

    if start_run_idx < 0 or end_run_idx < 0:
        return False

    start_run = paragraph.runs[start_run_idx]
    end_run = paragraph.runs[end_run_idx]

    start_run_global_start = positions[start_run_idx][1]
    end_run_global_start = positions[end_run_idx][1]

    prefix = start_run.text[: start - start_run_global_start]
    suffix = end_run.text[(end - end_run_global_start):]

    start_run.text = f"{prefix}{replacement}{suffix}"

    for idx in range(start_run_idx + 1, end_run_idx + 1):
        paragraph.runs[idx].text = ""

    return True


def _replace_placeholder_all_in_paragraph(paragraph, placeholder: str, replacement: str) -> None:
    while _replace_placeholder_once_in_paragraph(paragraph, placeholder, replacement):
        pass


def _replace_placeholders_everywhere(doc: Document, mapping: dict[str, str]) -> None:
    for paragraph in doc.paragraphs:
        for placeholder, replacement in mapping.items():
            _replace_placeholder_all_in_paragraph(paragraph, placeholder, replacement)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for placeholder, replacement in mapping.items():
                        _replace_placeholder_all_in_paragraph(paragraph, placeholder, replacement)


# ---------------------------------------------------------------------------
# Domain resolution helpers
# ---------------------------------------------------------------------------

def _resolve_proc_type(procurement: Any) -> str:
    materials = list(getattr(procurement, "materials", []) or [])
    is_services = any(bool(getattr(line, "is_service", False)) for line in materials)
    return "παροχή υπηρεσιών" if is_services else "προμήθεια υλικών"


def _resolve_handler_directory(procurement: Any) -> str:
    assignment = getattr(procurement, "handler_assignment", None)
    if assignment is not None:
        directory = getattr(assignment, "directory", None)
        if directory is not None:
            return _upper_no_accents(getattr(directory, "name", None))

    handler = getattr(procurement, "handler_personnel", None)
    directory = getattr(handler, "directory", None)
    return _upper_no_accents(getattr(directory, "name", None))


def _resolve_service_unit_place(service_unit: Any) -> str:
    _ = service_unit
    return "—"


def _participating_supplier_objects(procurement: Any) -> list[Any]:
    suppliers: list[Any] = []

    for link in list(getattr(procurement, "supplies_links", []) or []):
        supplier = getattr(link, "supplier", None)
        if supplier is not None:
            suppliers.append(supplier)

    return suppliers


def _economic_operator_label(procurement: Any) -> str:
    """
    - 1 supplier  -> στον οικονομικό φορέα
    - many        -> στους οικονομικούς φορείς
    """
    suppliers = _participating_supplier_objects(procurement)
    if len(suppliers) == 1:
        return "στον οικονομικό φορέα"
    return "στους οικονομικούς φορείς"


def _supplier_full_info(supplier: Any) -> str:
    """
    Requested format:
    NAME, ΑΦΜ: XXXXXXXXX, Διεύθυνση: ..., τηλέφωνο: ..., email
    """
    if supplier is None:
        return "—"

    parts: list[str] = []

    name = _safe(getattr(supplier, "name", None), default="")
    afm = _safe(getattr(supplier, "afm", None), default="")
    address = _safe(getattr(supplier, "address", None), default="")
    city = _safe(getattr(supplier, "city", None), default="")
    phone = _safe(getattr(supplier, "phone", None), default="")
    email = _safe(getattr(supplier, "email", None), default="")

    if name:
        parts.append(name)
    if afm:
        parts.append(f"ΑΦΜ: {afm}")

    address_parts: list[str] = []
    if address:
        address_parts.append(address)
    if city:
        address_parts.append(city)

    if address_parts:
        parts.append(f"Διεύθυνση: {', '.join(address_parts)}")

    if phone:
        parts.append(f"τηλέφωνο: {phone}")
    if email:
        parts.append(email)

    return ", ".join(parts).strip() or "—"


def _invited_suppliers_inline(procurement: Any) -> str:
    """
    Inline sentence version:
    «supplier1 ...», «supplier2 ...»
    """
    suppliers = _participating_supplier_objects(procurement)
    if not suppliers:
        return "—"

    return ", ".join(f"«{_supplier_full_info(supplier)}»" for supplier in suppliers)


def _recipients_block(procurement: Any) -> str:
    """
    Recipients section version:
    one supplier per line
    """
    suppliers = _participating_supplier_objects(procurement)
    if not suppliers:
        return "—"

    return "\n".join(f"«{_supplier_full_info(supplier)}»" for supplier in suppliers)


# ---------------------------------------------------------------------------
# Table helpers
# ---------------------------------------------------------------------------

def _set_cell_alignment(
    cell,
    *,
    horizontal: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.CENTER,
    vertical: WD_ALIGN_VERTICAL = WD_ALIGN_VERTICAL.CENTER,
) -> None:
    cell.vertical_alignment = vertical
    for paragraph in cell.paragraphs:
        paragraph.alignment = horizontal


def _find_invitation_table(doc: Document):
    """
    Find the main invitation materials/services table.

    Expected shape in current template:
    - 5 columns
    - header includes:
      Α/Α | ΠΕΡΙΓΡΑΦΗ | ΜΟΝ. ΜΕΤ. | ΣΥΝΟΛΙΚΗ ΠΟΣΟΤΗΤΑ | CPV
    """
    for table in doc.tables:
        if len(table.columns) != 5 or not table.rows:
            continue

        header = " ".join(cell.text.strip() for cell in table.rows[0].cells).upper()
        if (
            "Α/Α" in header
            and "ΠΕΡΙΓΡΑΦΗ" in header
            and "ΜΟΝ. ΜΕΤ." in header
            and "ΣΥΝΟΛΙΚΗ ΠΟΣΟΤΗΤΑ" in header
            and "CPV" in header
        ):
            return table

    return None


def _clear_table_body_keep_header(table, header_rows: int = 1) -> None:
    while len(table.rows) > header_rows:
        tbl = table._tbl
        tr = table.rows[header_rows]._tr
        tbl.remove(tr)


def _sorted_materials(procurement: Any) -> list[Any]:
    """
    Preserve source ordering by default.
    Change this function if you later want custom ordering.
    """
    return list(getattr(procurement, "materials", []) or [])


def _fill_invitation_table(table, procurement: Any) -> None:
    _clear_table_body_keep_header(table, header_rows=1)

    lines = _sorted_materials(procurement)

    if not lines:
        row = table.add_row().cells
        row[0].text = "—"
        row[1].text = "Δεν υπάρχουν γραμμές υλικών/υπηρεσιών."
        row[2].text = "—"
        row[3].text = "—"
        row[4].text = "—"
        for cell in row:
            _set_cell_alignment(cell)
        return

    for idx, line in enumerate(lines, start=1):
        row = table.add_row().cells
        row[0].text = str(idx)
        row[1].text = _safe(getattr(line, "description", None))
        row[2].text = _safe(getattr(line, "unit", None))

        qty_value = getattr(line, "quantity", None)
        row[3].text = _money_plain(qty_value) if qty_value is not None else "—"

        row[4].text = _safe(getattr(line, "cpv", None))

        _set_cell_alignment(row[0])
        _set_cell_alignment(row[1], horizontal=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_alignment(row[2])
        _set_cell_alignment(row[3])
        _set_cell_alignment(row[4])


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

@dataclass(frozen=True)
class InvitationConstants:
    pass


def build_invitation_docx(
    *,
    procurement: Any,
    service_unit: Any,
    winner: Optional[Any],
    analysis: dict[str, Any],
    constants: Optional[InvitationConstants] = None,
) -> bytes:
    _ = winner
    _ = analysis
    _ = constants

    template_path = _template_path()
    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")

    doc = Document(str(template_path))

    invited_suppliers_inline = _invited_suppliers_inline(procurement)
    recipients_info = _recipients_block(procurement)

    mapping: dict[str, str] = {
        "{{SERVICE_UNIT_NAME}}": _upper_no_accents(getattr(service_unit, "description", None)),
        "{{SERVICE_UNIT_DESCRIPTION}}": _safe(getattr(service_unit, "description", None)),
        "{{PROCUREMENT_SHORT_DESCRIPTION}}": _lower_preserve_accents(
            getattr(procurement, "description", None)
        ),
        "{{HANDLER_DIRECTORY}}": _resolve_handler_directory(procurement),
        "{{SERVICE_UNIT_PHONE}}": _safe(getattr(service_unit, "phone", None)),
        "{{ SERVICE_UNIT_PHONE}}": _safe(getattr(service_unit, "phone", None)),
        "{{SERVICE_UNIT_REGION}}": _safe(getattr(service_unit, "region", None)),
        "{{SHORT_DATE}}": _short_date_el(),
        "{{PROC_TYPE}}": _resolve_proc_type(procurement),
        "{{ECONOMIC_OPERATOR_LABEL}}": _economic_operator_label(procurement),
        "{{procurement.aay}}": _safe(getattr(procurement, "aay", None)),
        "{{procurement.adam_aay}}": _safe(getattr(procurement, "adam_aay", None)),
        "{{procurement hop_approval_commitment}}": _safe(
            getattr(procurement, "hop_approval_commitment", None)
        ),
        "{{INVITED_SUPPLIERS_INLINE}}": invited_suppliers_inline,
        "{{RECIPIENTS_INFO}}": recipients_info,

        # Backward compatibility
        "{{WINNER_SUPPLIER_LINE}}": invited_suppliers_inline,

        "{{SERVICE_UNIT_ADRESS}}": _safe(getattr(service_unit, "address", None)),
        "{{ SERVICE_UNIT_ADRESS}}": _safe(getattr(service_unit, "address", None)),
        "{{SERVICE_UNIT_PLACE}}": _resolve_service_unit_place(service_unit),
        "{{SERVICE_UNIT_POSTAL_CODE}}": _safe(getattr(service_unit, "postal_code", None)),
        "{{ SERVICE_UNIT_EMAIL}}": _safe(getattr(service_unit, "email", None)),
        "{{SERVICE_UNIT_EMAIL}}": _safe(getattr(service_unit, "email", None)),
        "{{service.commander}}": _safe(getattr(service_unit, "commander", None)),
        "{{COMMANDER_ROLE_TYPE}}": _safe(getattr(service_unit, "commander_role_type", None)),

        # Keep old material placeholders empty for backward compatibility,
        # because the table is now populated with real rows.
        "{{ML_NO}}": "",
        "{{ML_DESC}}": "",
        "{{ML_UNIT}}": "",
        "{{ML_QTY}}": "",
        "{{ML_CPV}}": "",
    }

    _replace_placeholders_everywhere(doc, mapping)

    invitation_table = _find_invitation_table(doc)
    if invitation_table is not None:
        _fill_invitation_table(invitation_table, procurement)

    _set_global_font_arial_12(doc)

    output = BytesIO()
    doc.save(output)
    return output.getvalue()


def build_invitation_filename(
    *,
    procurement: Any,
    winner: Optional[Any],
) -> str:
    _ = winner

    description = _safe(
        getattr(procurement, "description", None),
        default=f"Προμήθεια #{getattr(procurement, 'id', '—')}",
    )
    return f"Πρόσκληση Υποβολής Προσφοράς για την {description}.docx"