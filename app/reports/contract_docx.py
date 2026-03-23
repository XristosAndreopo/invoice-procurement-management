"""
app/reports/contract_docx.py

Generate "ΣΥΜΒΑΣΗ" as DOCX bytes using the provided Word template
`contact_template.docx`.

SOURCE OF TRUTH
---------------
This implementation is aligned to:
- combined_project.md
- current uploaded `contact_template.docx`
- current Procurement / ServiceUnit / Supplier / Organization models
- current payment-analysis service

IMPORTANT BUSINESS RULE
-----------------------
The document switches wording dynamically between:
- υπηρεσίες
- υλικά

Canonical project rule:
- if any procurement material line has `is_service == True`
  -> services contract
- otherwise
  -> goods/materials contract

TEMPLATE CONTRACT
-----------------
This implementation supports:
1. explicit placeholders such as:
   - {{CONTRACT_KIND_TITLE}}
   - {{TABLE_REFERENCE_PHRASE}}
   - {{DELIVERY_SENTENCE}}
   - {{DEADLINE_PHRASE}}
   - {{COMPLETION_PHRASE}}
   - {{NONCONFORMITY_SUBJECT}}
   - {{TABLE_TITLE}}
   - {{VAT_SENTENCE}}
   - {{PROC_TYPE_LOWER}}
   - {{DELIVERY_TO_SERVICE_SENTENCE}}

2. legacy / typo placeholders kept for backward compatibility, such as:
   - {{WINNER_SUPPLIER_DESCREIPTION}}
   - {{PROCUREMENT_INDENTITY_PROSKLISIS}}
   - {{PROCUREMENT_INDENTITY_APOFASIS_ANATHESIS}}

KNOWN MODEL LIMITATIONS
-----------------------
The current source-of-truth models do NOT provide dedicated fields for:
- supplier legal representative name
- representative ID card number
- dedicated contract signing place
- explicit contract delivery deadline date

Therefore:
- {{SERVICE_UNIT_PLACE}} resolves conservatively to:
  service_unit.region -> service_unit.prefecture -> "—"
- legal-representative free text remains template-owned if present outside
  placeholders
"""

from __future__ import annotations

import unicodedata
from dataclasses import dataclass
from datetime import date, datetime
from decimal import Decimal
from io import BytesIO
from pathlib import Path
from typing import Any, Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


# ---------------------------------------------------------------------------
# Generic helpers
# ---------------------------------------------------------------------------

def _safe(value: Any, default: str = "—") -> str:
    text = ("" if value is None else str(value)).strip()
    return text if text else default


def _to_decimal(value: Any) -> Decimal:
    try:
        return Decimal(str(value or "0"))
    except Exception:
        return Decimal("0.00")


def _money_plain(value: Any) -> str:
    amount = _to_decimal(value).quantize(Decimal("0.01"))
    text = f"{amount:,.2f}"
    return text.replace(",", "X").replace(".", ",").replace("X", ".")


def _money(value: Any) -> str:
    return f"{_money_plain(value)} €"


def _percent(value: Any) -> str:
    amount = _to_decimal(value).quantize(Decimal("0.01"))
    text = f"{amount:,.2f}"
    return text.replace(",", "X").replace(".", ",").replace("X", ".")


def _upper_no_accents(value: Any, default: str = "—") -> str:
    text = _safe(value, default=default)
    normalized = unicodedata.normalize("NFD", text)
    no_marks = "".join(ch for ch in normalized if not unicodedata.combining(ch))
    return unicodedata.normalize("NFC", no_marks).upper()


def _upper_service_name(value: Any) -> str:
    return _upper_no_accents(value)


def _format_date(value: Any, default: str = "—") -> str:
    if value is None:
        return default

    if isinstance(value, date):
        return value.strftime("%d/%m/%Y")

    try:
        return value.strftime("%d/%m/%Y")
    except Exception:
        text = str(value).strip()
        return text if text else default


def _short_date_el(value: Any | None = None) -> str:
    months = {
        1: "Ιαν",
        2: "Φεβ",
        3: "Μαρ",
        4: "Απρ",
        5: "Μαϊ",
        6: "Ιουν",
        7: "Ιουλ",
        8: "Αυγ",
        9: "Σεπ",
        10: "Οκτ",
        11: "Νοε",
        12: "Δεκ",
    }

    dt = value or datetime.now()

    try:
        day = int(dt.day)
        month = int(dt.month)
        year_2d = int(dt.year) % 100
    except Exception:
        dt = datetime.now()
        day = dt.day
        month = dt.month
        year_2d = dt.year % 100

    month_label = months.get(month, "")
    return f"{day:02d} {month_label} {year_2d:02d}".strip()


# ---------------------------------------------------------------------------
# Amount-to-words helpers
# ---------------------------------------------------------------------------

def _int_to_greek_words_genitive(n: int) -> str:
    if n < 0:
        raise ValueError("Negative values are not supported.")
    if n == 0:
        return "μηδενός"

    units = {
        0: "",
        1: "ενός",
        2: "δύο",
        3: "τριών",
        4: "τεσσάρων",
        5: "πέντε",
        6: "έξι",
        7: "επτά",
        8: "οκτώ",
        9: "εννέα",
    }

    teens = {
        10: "δέκα",
        11: "έντεκα",
        12: "δώδεκα",
        13: "δεκατριών",
        14: "δεκατεσσάρων",
        15: "δεκαπέντε",
        16: "δεκαέξι",
        17: "δεκαεπτά",
        18: "δεκαοκτώ",
        19: "δεκαεννέα",
    }

    tens = {
        2: "είκοσι",
        3: "τριάντα",
        4: "σαράντα",
        5: "πενήντα",
        6: "εξήντα",
        7: "εβδομήντα",
        8: "ογδόντα",
        9: "ενενήντα",
    }

    hundreds = {
        1: "εκατόν",
        2: "διακοσίων",
        3: "τριακοσίων",
        4: "τετρακοσίων",
        5: "πεντακοσίων",
        6: "εξακοσίων",
        7: "επτακοσίων",
        8: "οκτακοσίων",
        9: "εννιακοσίων",
    }

    def two_digits(num: int) -> str:
        if num < 10:
            return units[num]
        if 10 <= num <= 19:
            return teens[num]

        t = num // 10
        u = num % 10
        if u == 0:
            return tens[t]
        return f"{tens[t]} {units[u]}".strip()

    def three_digits(num: int) -> str:
        if num < 100:
            return two_digits(num)

        h = num // 100
        rem = num % 100

        if rem == 0:
            if h == 1:
                return "εκατό"
            return hundreds[h]

        return f"{hundreds[h]} {two_digits(rem)}".strip()

    parts: list[str] = []

    millions = n // 1_000_000
    remainder = n % 1_000_000

    thousands = remainder // 1_000
    below_thousand = remainder % 1_000

    if millions:
        if millions == 1:
            parts.append("ενός εκατομμυρίου")
        else:
            parts.append(f"{three_digits(millions)} εκατομμυρίων")

    if thousands:
        if thousands == 1:
            parts.append("χιλίων")
        else:
            parts.append(f"{three_digits(thousands)} χιλιάδων")

    if below_thousand:
        parts.append(three_digits(below_thousand))

    return " ".join(p for p in parts if p).strip()


def _money_words_el(value: Any) -> str:
    amount = _to_decimal(value).quantize(Decimal("0.01"))

    euros = int(amount)
    cents = int((amount - Decimal(euros)) * 100)

    euro_words = _int_to_greek_words_genitive(euros)

    if cents == 0:
        return f"{euro_words} ευρώ"

    cents_words = _int_to_greek_words_genitive(cents)
    return f"{euro_words} ευρώ και {cents_words} λεπτών"


# ---------------------------------------------------------------------------
# Template helpers
# ---------------------------------------------------------------------------

def _template_path() -> Path:
    app_dir = Path(__file__).resolve().parents[1]
    return app_dir / "templates" / "docx" / "contact_template.docx"


def _set_global_font_arial_12(doc: Document) -> None:
    try:
        style = doc.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(12)
    except Exception:
        pass


def _paragraph_runs_text(paragraph) -> tuple[str, list[tuple[int, int, int]]]:
    pieces: list[str] = []
    positions: list[tuple[int, int, int]] = []

    cursor = 0
    for idx, run in enumerate(paragraph.runs):
        text = run.text or ""
        pieces.append(text)
        start = cursor
        end = start + len(text)
        positions.append((idx, start, end))
        cursor = end

    return "".join(pieces), positions


def _find_run_index_at_offset(positions: list[tuple[int, int, int]], offset: int) -> int:
    for run_index, start, end in positions:
        if start <= offset < end:
            return run_index

    if positions and offset == positions[-1][2]:
        return positions[-1][0]

    return -1


def _replace_placeholder_once_in_paragraph(paragraph, placeholder: str, replacement: str) -> bool:
    full_text, positions = _paragraph_runs_text(paragraph)
    if not full_text or placeholder not in full_text:
        return False

    start = full_text.find(placeholder)
    end = start + len(placeholder)

    start_run_idx = _find_run_index_at_offset(positions, start)
    end_run_idx = _find_run_index_at_offset(positions, end - 1)

    if start_run_idx < 0 or end_run_idx < 0:
        return False

    start_run = paragraph.runs[start_run_idx]
    end_run = paragraph.runs[end_run_idx]

    start_run_global_start = positions[start_run_idx][1]
    end_run_global_start = positions[end_run_idx][1]

    prefix = start_run.text[: start - start_run_global_start]
    suffix = end_run.text[(end - end_run_global_start):]

    start_run.text = f"{prefix}{replacement}{suffix}"

    for idx in range(start_run_idx + 1, end_run_idx + 1):
        paragraph.runs[idx].text = ""

    return True


def _replace_placeholder_all_in_paragraph(paragraph, placeholder: str, replacement: str) -> None:
    while _replace_placeholder_once_in_paragraph(paragraph, placeholder, replacement):
        pass


def _replace_placeholders_everywhere(doc: Document, mapping: dict[str, str]) -> None:
    for paragraph in doc.paragraphs:
        for placeholder, replacement in mapping.items():
            _replace_placeholder_all_in_paragraph(paragraph, placeholder, replacement)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for placeholder, replacement in mapping.items():
                        _replace_placeholder_all_in_paragraph(paragraph, placeholder, replacement)


def _replace_placeholders_in_headers_and_footers(doc: Document, mapping: dict[str, str]) -> None:
    """
    Replace placeholders in all header/footer variants for every section:
    - default header/footer
    - first-page header/footer
    - even-page header/footer

    This covers Word documents that use:
    - Different First Page
    - Different Odd & Even Pages
    """
    for section in doc.sections:
        containers = [
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ]

        for container in containers:
            if container is None:
                continue

            for paragraph in container.paragraphs:
                for placeholder, replacement in mapping.items():
                    _replace_placeholder_all_in_paragraph(paragraph, placeholder, replacement)

            for table in container.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for placeholder, replacement in mapping.items():
                                _replace_placeholder_all_in_paragraph(paragraph, placeholder, replacement)


def _replace_literal_text_everywhere(doc: Document, replacements: dict[str, str]) -> None:
    """
    Replace literal fallback text in paragraphs and table cells.

    This stays as a defensive fallback for older template versions that still
    contain slash-based text instead of explicit placeholders.
    """
    for paragraph in doc.paragraphs:
        original = paragraph.text or ""
        updated = original
        for src, dst in replacements.items():
            if src in updated:
                updated = updated.replace(src, dst)
        if updated != original:
            paragraph.text = updated

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    original = paragraph.text or ""
                    updated = original
                    for src, dst in replacements.items():
                        if src in updated:
                            updated = updated.replace(src, dst)
                    if updated != original:
                        paragraph.text = updated


# ---------------------------------------------------------------------------
# Domain resolution helpers
# ---------------------------------------------------------------------------

def _resolve_is_services(procurement: Any) -> bool:
    materials = list(getattr(procurement, "materials", []) or [])
    return any(bool(getattr(line, "is_service", False)) for line in materials)


def _resolve_proc_type(is_services: bool) -> str:
    return "ΠΑΡΟΧΗΣ ΥΠΗΡΕΣΙΩΝ" if is_services else "ΠΡΟΜΗΘΕΙΑΣ ΥΛΙΚΩΝ"


def _resolve_proc_type_lower(is_services: bool) -> str:
    return "παροχή υπηρεσιών" if is_services else "προμήθεια υλικών"


def _resolve_contract_kind_title(is_services: bool) -> str:
    return "Παροχής Υπηρεσιών" if is_services else "Προμήθειας Υλικών"


def _resolve_table_title(is_services: bool) -> str:
    return (
        "ΠΙΝΑΚΑΣ ΠΑΡΟΧΗΣ ΥΠΗΡΕΣΙΩΝ"
        if is_services
        else "ΠΙΝΑΚΑΣ ΥΠΟ ΠΡΟΜΗΘΕΙΑ ΥΛΙΚΩΝ"
    )


def _resolve_table_reference_phrase(is_services: bool) -> str:
    return (
        "πίνακα παρεχόμενων υπηρεσιών"
        if is_services
        else "πίνακα υπό προμήθεια υλικών"
    )


def _resolve_delivery_sentence(is_services: bool) -> str:
    return (
        "Οι υπηρεσίες θα παρασχεθούν"
        if is_services
        else "Τα υπό προμήθεια υλικά θα παραδοθούν"
    )


def _resolve_delivery_to_service_sentence(is_services: bool) -> str:
    return (
        "Οι υπηρεσίες θα εκτελεστούν"
        if is_services
        else "Τα υλικά θα παραδοθούν στην υπηρεσία"
    )


def _resolve_deadline_phrase(is_services: bool) -> str:
    return "εκτέλεσης εργασιών" if is_services else "παράδοσης των υλικών"


def _resolve_completion_phrase(is_services: bool) -> str:
    return "εκτέλεσης εργασιών" if is_services else "παράδοσης υλικών"


def _resolve_nonconformity_subject(is_services: bool) -> str:
    return "οι παρεχόμενες υπηρεσίες" if is_services else "τα υπό προμήθεια υλικά"


def _resolve_item_label_plural(is_services: bool) -> str:
    return "υπηρεσιών" if is_services else "υλικών"


def _resolve_vat_sentence(analysis: dict[str, Any], is_services: bool) -> str:
    vat_percent = _to_decimal(analysis.get("vat_percent", 0)).quantize(Decimal("0.01"))
    subject = "παρεχόμενων υπηρεσιών" if is_services else "υπό προμήθεια υλικών"

    if vat_percent == Decimal("0.00"):
        return f"Η τιμή των {subject} απαλλάσσεται Φόρου Προστιθέμενης Αξίας (Φ.Π.Α.)."

    return (
        f"Η τιμή των {subject} επιβαρύνεται με {_percent(vat_percent)}% "
        "Φόρου Προστιθέμενης Αξίας (Φ.Π.Α.)."
    )


def _winner_supplier_name(winner: Any) -> str:
    return _safe(getattr(winner, "name", None), default="—")


def _winner_supplier_afm(winner: Any) -> str:
    return _safe(getattr(winner, "afm", None), default="—")


def _winner_supplier_emba(winner: Any) -> str:
    return _safe(getattr(winner, "emba", None), default="—")


def _resolve_handler_directory(procurement: Any) -> str:
    assignment = getattr(procurement, "handler_assignment", None)
    if assignment is not None:
        directory = getattr(assignment, "directory", None)
        if directory is not None:
            return _upper_no_accents(getattr(directory, "name", None))

    handler = getattr(procurement, "handler_personnel", None)
    directory = getattr(handler, "directory", None)
    return _upper_no_accents(getattr(directory, "name", None))


def _resolve_service_unit_place(service_unit: Any) -> str:
    region = getattr(service_unit, "region", None)
    if region:
        return _safe(region)

    prefecture = getattr(service_unit, "prefecture", None)
    if prefecture:
        return _safe(prefecture)

    return "—"


def _resolve_document_total_value(procurement: Any, analysis: dict[str, Any]) -> Any:
    grand_total = getattr(procurement, "grand_total", None)
    if grand_total is not None:
        return grand_total

    payable_total = analysis.get("payable_total")
    if payable_total is not None:
        return payable_total

    return analysis.get("sum_total", 0)


def _resolve_document_total(procurement: Any, analysis: dict[str, Any]) -> str:
    return _money_plain(_resolve_document_total_value(procurement, analysis))


def _resolve_krathseis_label(analysis: dict[str, Any]) -> str:
    public_withholdings = analysis.get("public_withholdings") or {}
    return f"{_percent(public_withholdings.get('total_percent', 0))}%"


def _resolve_fe_label(analysis: dict[str, Any]) -> str:
    income_tax = analysis.get("income_tax") or {}
    return _percent(income_tax.get("rate_percent", 0))


def _resolve_fpa_label(analysis: dict[str, Any]) -> str:
    return f"{_percent(analysis.get('vat_percent', 0))}%"


# ---------------------------------------------------------------------------
# Table helpers
# ---------------------------------------------------------------------------

def _set_cell_alignment(
    cell,
    *,
    horizontal: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.CENTER,
    vertical: WD_ALIGN_VERTICAL = WD_ALIGN_VERTICAL.CENTER,
) -> None:
    cell.vertical_alignment = vertical
    for paragraph in cell.paragraphs:
        paragraph.alignment = horizontal


def _find_contract_table(doc: Document):
    """
    Find the main contract pricing table.

    Current expected shape:
    - 6 columns
    - header includes: ΠΕΡΙΓΡΑΦΗ, ΤΙΜΗ, ΣΥΝΟΛΟ
    """
    for table in doc.tables:
        if len(table.columns) != 6 or not table.rows:
            continue

        header = " ".join(cell.text for cell in table.rows[0].cells).upper()
        if "ΠΕΡΙΓΡΑΦΗ" in header and "ΤΙΜΗ" in header and "ΣΥΝΟΛΟ" in header:
            return table

    return None


def _clear_table_body_keep_header(table, header_rows: int = 1) -> None:
    while len(table.rows) > header_rows:
        tbl = table._tbl
        tr = table.rows[header_rows]._tr
        tbl.remove(tr)


def _add_summary_row(table, label: str, amount: Any) -> None:
    row = table.add_row().cells
    merged = row[0].merge(row[4])
    merged.text = label
    row[5].text = _money_plain(amount)

    _set_cell_alignment(
        merged,
        horizontal=WD_ALIGN_PARAGRAPH.RIGHT,
        vertical=WD_ALIGN_VERTICAL.CENTER,
    )
    _set_cell_alignment(
        row[5],
        horizontal=WD_ALIGN_PARAGRAPH.CENTER,
        vertical=WD_ALIGN_VERTICAL.CENTER,
    )


def _fill_contract_table(table, procurement: Any, analysis: dict[str, Any]) -> None:
    _clear_table_body_keep_header(table, header_rows=1)

    lines = list(getattr(procurement, "materials", []) or [])

    if not lines:
        row = table.add_row().cells
        row[0].text = "—"
        row[1].text = "Δεν υπάρχουν γραμμές υλικών/υπηρεσιών."
        row[2].text = "—"
        row[3].text = "—"
        row[4].text = "—"
        row[5].text = "—"
        for cell in row:
            _set_cell_alignment(cell)
        return

    for idx, line in enumerate(lines, start=1):
        row = table.add_row().cells
        row[0].text = str(idx)
        row[1].text = _safe(getattr(line, "description", None), default="")
        row[2].text = _safe(getattr(line, "unit", None))
        row[3].text = _safe(getattr(line, "quantity", None))
        row[4].text = _money_plain(getattr(line, "unit_price", None))
        row[5].text = _money_plain(getattr(line, "total_pre_vat", None))

        for cell in row:
            _set_cell_alignment(cell)

    public_withholdings = analysis.get("public_withholdings") or {}
    income_tax = analysis.get("income_tax") or {}

    _add_summary_row(table, "Μερικό Σύνολο", analysis.get("sum_total", 0))
    _add_summary_row(
        table,
        f"Κρατήσεις Υπέρ Δημοσίου ({_percent(public_withholdings.get('total_percent', 0))}%)",
        public_withholdings.get("total_amount", 0),
    )
    _add_summary_row(
        table,
        f"Φόρος Εισοδήματος ({_percent(income_tax.get('rate_percent', 0))}%)",
        income_tax.get("amount", 0),
    )
    _add_summary_row(
        table,
        f"ΦΠΑ ({_percent(analysis.get('vat_percent', 0))}%)",
        analysis.get("vat_amount", 0),
    )
    _add_summary_row(table, "Τελικό Σύνολο", analysis.get("payable_total", 0))


# ---------------------------------------------------------------------------
# Legacy slash-text fallback
# ---------------------------------------------------------------------------

def _apply_legacy_goods_services_wording(
    doc: Document,
    *,
    is_services: bool,
    analysis: dict[str, Any],
) -> None:
    """
    Defensive fallback for older template versions that still contain slash
    wording instead of explicit placeholders.
    """
    title_value = _resolve_contract_kind_title(is_services)
    table_title = _resolve_table_title(is_services)
    table_ref = _resolve_table_reference_phrase(is_services)
    delivery_sentence = _resolve_delivery_sentence(is_services)
    deadline_phrase = _resolve_deadline_phrase(is_services)
    completion_phrase = _resolve_completion_phrase(is_services)
    subject_value = _resolve_nonconformity_subject(is_services)
    item_label_plural = _resolve_item_label_plural(is_services)
    vat_sentence = _resolve_vat_sentence(analysis, is_services)
    proc_type_lower = _resolve_proc_type_lower(is_services)
    delivery_to_service = _resolve_delivery_to_service_sentence(is_services)

    replacements = {
        "Παροχής Υπηρεσιών/Προμήθειας Υλικών": title_value,
        "πίνακα παρεχόμενων υπηρεσιών/υπό προμήθεια υλικών": table_ref,
        "πίνακα παρεχόμενων υπηρεσιών/προμηθευτέων υλικών": table_ref,
        "Οι υπηρεσίες θα παρασχεθούν/Τα υπό προμήθεια υλικά θα παραδοθούν": delivery_sentence,
        "Τα υλικά θα παραδοθούν στην υπηρεσία /Οι υπηρεσίες θα εκτελεστούν": delivery_to_service,
        "Τα υλικά θα παραδοθούν στην υπηρεσία/Οι υπηρεσίες θα εκτελεστούν": delivery_to_service,
        "παράδοσης των υλικών/εκτέλεσης εργασιών": deadline_phrase,
        "παράδοσης υλικών/εκτέλεσης εργασιών": completion_phrase,
        "τα υπό προμήθεια υλικά/οι παρεχόμενες υπηρεσίες": subject_value,
        "ΠΙΝΑΚΑΣ ΠΑΡΟΧΗΣ ΥΠΗΡΕΣΙΩΝ/ΥΠΟ ΠΡΟΜΗΘΕΙΑ ΥΛΙΚΩΝ": table_title,
        "ΠΡΟΜΗΘΕΙΑΣ ΥΛΙΚΩΝ/ΠΑΡΟΧΗΣ ΥΠΗΡΕΣΙΩΝ": _resolve_proc_type(is_services),
        "προμήθειας υλικών/παροχής υπηρεσιών": proc_type_lower,
        "προμήθειας υλικών / παροχής υπηρεσιών": proc_type_lower,
        "Η τιμή των παρεχόμενων υπηρεσιών/προμηθευτέων υλικών απαλλάσσεται Φόρου Προστιθέμενης Αξίας (Φ.Π.Α.)/ επιβαρύνεται με {{PROCUREMENT_FPA}} Φόρου Προστιθέμενης Αξίας (Φ.Π.Α.).":
            vat_sentence,
        "Η τιμή των παρεχόμενων υπηρεσιών/προμηθυτέων υλικών απαλλάσσεται Φόρου Προστιθέμενης Αξίας (Φ.Π.Α.)/ επιβαρύνεται με {{PROCUREMENT_FPA}} Φόρου Προστιθέμενης Αξίας (Φ.Π.Α.).":
            vat_sentence,
        "Επί της καθαρής αξίας των υπηρεσιών υπολογίζονται κρατήσεις":
            f"Επί της καθαρής αξίας των {item_label_plural} υπολογίζονται κρατήσεις",
    }

    _replace_literal_text_everywhere(doc, replacements)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

@dataclass(frozen=True)
class ContractConstants:
    """
    Future-proof constants container for contract report generation.
    """
    pass


def build_contract_docx(
    *,
    procurement: Any,
    service_unit: Any,
    winner: Optional[Any],
    analysis: dict[str, Any],
    constants: Optional[ContractConstants] = None,
) -> bytes:
    _ = constants

    template_path = _template_path()
    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")

    doc = Document(str(template_path))

    is_services = _resolve_is_services(procurement)
    proc_type = _resolve_proc_type(is_services)

    mapping: dict[str, str] = {
        # Core document placeholders
        "{{SERVICE_UNIT_NAME}}": _upper_service_name(
            _safe(getattr(service_unit, "description", None), default="—")
        ),
        "{{HANDLER_DIRECTORY}}": _resolve_handler_directory(procurement),
        "{{SERVICE_UNIT_PHONE}}": _safe(getattr(service_unit, "phone", None)),
        "{{PROC_TYPE}}": proc_type,
        "{{PROC_TYPE_LOWER}}": _resolve_proc_type_lower(is_services),
        "{{PROCUREMENT_CONTRACT_NUMBER}}": _safe(getattr(procurement, "contract_number", None)),
        "{{WINNER_SUPPLIER_AFM}}": _winner_supplier_afm(winner),
        "{{ WINNER_SUPPLIER_AFM}}": _winner_supplier_afm(winner),
        "{{WINNER_SUPPLIER_DESCRIPTION}}": _winner_supplier_name(winner),
        "{{WINNER_SUPPLIER_DESCREIPTION}}": _winner_supplier_name(winner),  # legacy typo
        "{{WINNER_SUPPLIER_EMBA}}": _winner_supplier_emba(winner),
        "{{PROCUREMENT_HOP_APPROVAL}}": _safe(getattr(procurement, "hop_approval", None)),
        "{{PROCUREMENT_AAY}}": _safe(getattr(procurement, "aay", None)),
        "{{PROCUREMENT_ADAM_AAY}}": _safe(getattr(procurement, "adam_aay", None)),
        "{{PROCUREMENT_ADA_AAY}}": _safe(getattr(procurement, "ada_aay", None)),

        # Identity placeholders - corrected and legacy variants
        "{{PROCUREMENT_IDENTITY_PROSKLISIS}}": _safe(
            getattr(procurement, "identity_prosklisis", None)
        ),
        "{{PROCUREMENT_INDENTITY_PROSKLISIS}}": _safe(
            getattr(procurement, "identity_prosklisis", None)
        ),
        "{{PROCUREMENT_ADAM_PROSKLISIS}}": _safe(
            getattr(procurement, "adam_prosklisis", None)
        ),
        "{{ PROCUREMENT_ADAM_PROSKLISIS}}": _safe(
            getattr(procurement, "adam_prosklisis", None)
        ),

        "{{PROCUREMENT_IDENTITY_APOFASIS_ANATHESIS}}": _safe(
            getattr(procurement, "identity_apofasis_anathesis", None)
        ),
        "{{PROCUREMENT_INDENTITY_APOFASIS_ANATHESIS}}": _safe(
            getattr(procurement, "identity_apofasis_anathesis", None)
        ),
        "{{PROCUREMENT_ADAM_APOFASIS_ANATHESIS}}": _safe(
            getattr(procurement, "adam_apofasis_anathesis", None)
        ),

        # Service / commander / accounting
        "{{SERVICE_UNIT_PLACE}}": _resolve_service_unit_place(service_unit),
        "{{COMMANDER_ROLE_TYPE}}": _safe(getattr(service_unit, "commander_role_type", None)),
        "{{SERVICE_COMMANDER}}": _safe(getattr(service_unit, "commander", None)),
        "{{SERVICE_UNIT_COMMANDER}}": _safe(getattr(service_unit, "commander", None)),
        "{{SERVICE.AAHT}}": _safe(getattr(service_unit, "aahit", None)),
        "{{PROCUREMENT_ALE}}": _safe(getattr(procurement, "ale", None)),
        "{{CURRENT_YEAR}}": str(getattr(procurement, "fiscal_year", None) or datetime.now().year),

        # Financial placeholders
        "{{ML_TOTAL_WORDS}}": _money_words_el(_resolve_document_total_value(procurement, analysis)),
        "{{ML_TOTAL}}": _resolve_document_total(procurement, analysis),
        "{{PROCUREMENT_KRATHSEIS_SYNOLO}}": _resolve_krathseis_label(analysis),
        "{{PROCUREMENT_FE}}": _resolve_fe_label(analysis),
        "{{PROCUREMENT_FPA}}": _resolve_fpa_label(analysis),
        "{{AN_SUM_TOTAL}}": _money_plain(analysis.get("sum_total", 0)),
        "{{VAT_SENTENCE}}": _resolve_vat_sentence(analysis, is_services),

        # New explicit dynamic placeholders
        "{{CONTRACT_KIND_TITLE}}": _resolve_contract_kind_title(is_services),
        "{{TABLE_REFERENCE_PHRASE}}": _resolve_table_reference_phrase(is_services),
        "{{DELIVERY_SENTENCE}}": _resolve_delivery_sentence(is_services),
        "{{DELIVERY_TO_SERVICE_SENTENCE}}": _resolve_delivery_to_service_sentence(is_services),
        "{{DEADLINE_PHRASE}}": _resolve_deadline_phrase(is_services),
        "{{COMPLETION_PHRASE}}": _resolve_completion_phrase(is_services),
        "{{NONCONFORMITY_SUBJECT}}": _resolve_nonconformity_subject(is_services),
        "{{TABLE_TITLE}}": _resolve_table_title(is_services),
    }

    _replace_placeholders_everywhere(doc, mapping)
    _replace_placeholders_in_headers_and_footers(doc, mapping)

    # Keep legacy fallback for older template copies.
    _apply_legacy_goods_services_wording(doc, is_services=is_services, analysis=analysis)

    contract_table = _find_contract_table(doc)
    if contract_table is not None:
        _fill_contract_table(contract_table, procurement, analysis)

    _set_global_font_arial_12(doc)

    output = BytesIO()
    doc.save(output)
    return output.getvalue()


def build_contract_filename(
    *,
    procurement: Any,
    winner: Optional[Any],
    is_services: bool,
) -> str:
    kind = "Παροχής Υπηρεσιών" if is_services else "Προμήθειας Υλικών"
    supplier_name = _safe(getattr(winner, "name", None), default="—")

    grand_total = getattr(procurement, "grand_total", None)
    if grand_total is None and hasattr(procurement, "compute_payment_analysis"):
        try:
            computed_analysis = procurement.compute_payment_analysis()
            grand_total = (
                computed_analysis.get("payable_total")
                or computed_analysis.get("grand_total")
                or computed_analysis.get("sum_total")
            )
        except Exception:
            grand_total = None

    total_str = _money_plain(grand_total)
    return f"Σύμβαση {kind} {supplier_name} {total_str}.docx"