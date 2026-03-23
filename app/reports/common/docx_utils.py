"""
Shared low-level DOCX helpers for report generation.

This module centralizes:
- placeholder replacement across split runs
- header/footer replacement
- paragraph insertion/cloning
- common font normalization
- table cleanup / alignment helpers

PERFORMANCE NOTES
-----------------
The DOCX report builders rely heavily on placeholder replacement and, after that
was optimized, font normalization became one of the main remaining shared costs.

The optimized implementation below keeps the existing public API intact while
reducing repeated work:

Placeholder replacement:
- paragraphs without '{{' are skipped immediately
- each candidate paragraph is flattened only once
- placeholder matches are discovered from the paragraph text itself
- all replacements for a paragraph are applied in a single rebuild pass
- split-run placeholder support is preserved

Font normalization:
- runs already set to Arial 12 are skipped
- style writes are only performed when needed
- traversal coverage remains unchanged

IMPORTANT COMPATIBILITY RULE
----------------------------
This file is intentionally implemented as a low-risk performance patch:
- function names stay the same
- signatures stay the same
- call sites do not need to change
"""

from __future__ import annotations

import re
from copy import deepcopy
from typing import Iterable

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Pt

# Shared constant target font configuration.
_TARGET_FONT_NAME = "Arial"
_TARGET_FONT_SIZE_PT = 12
_TARGET_FONT_SIZE = Pt(_TARGET_FONT_SIZE_PT)

# Fast sentinel used to skip paragraphs/cells that obviously do not contain
# placeholders.
_PLACEHOLDER_SENTINEL = "{{"

# Matches literal placeholder tokens such as:
#   {{FIELD_NAME}}
#
# It intentionally does not try to parse nested braces.
_PLACEHOLDER_TOKEN_RE = re.compile(r"\{\{[^{}]+\}\}")


def _font_name_matches(run, expected_name: str) -> bool:
    """
    Return True when the run already uses the expected font name.

    `python-docx` may expose font name as None or as a concrete string depending
    on direct formatting vs inherited styling. We only skip a write when the
    concrete current value is already the desired one.
    """
    try:
        return run.font.name == expected_name
    except Exception:
        return False


def _font_size_matches_pt(run, expected_pt: int) -> bool:
    """
    Return True when the run already uses the expected point size.

    `run.font.size` is typically an EMU-like Length object. Comparing its `.pt`
    value is the most stable way to avoid unnecessary writes while preserving
    current behavior.
    """
    try:
        size = run.font.size
        if size is None:
            return False
        return round(float(size.pt), 2) == float(expected_pt)
    except Exception:
        return False


def _normalize_run_font_if_needed(run) -> None:
    """
    Normalize one run to Arial 12, avoiding unnecessary property writes.

    This preserves the existing normalization behavior while reducing repeated
    writes for runs that are already in the desired state.
    """
    if not _font_name_matches(run, _TARGET_FONT_NAME):
        run.font.name = _TARGET_FONT_NAME

    if not _font_size_matches_pt(run, _TARGET_FONT_SIZE_PT):
        run.font.size = _TARGET_FONT_SIZE


def _normalize_paragraph_runs_font(paragraph) -> None:
    """
    Normalize every run in a paragraph to Arial 12.
    """
    for run in paragraph.runs:
        _normalize_run_font_if_needed(run)


def _normalize_table_runs_font(table) -> None:
    """
    Normalize every run inside a table's cells to Arial 12.
    """
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _normalize_paragraph_runs_font(paragraph)


def set_global_font_arial_12(doc: Document) -> None:
    """
    Normalize generated document font to Arial 12 where possible.

    This keeps current behavior consistent with the existing report modules.

    PERFORMANCE
    -----------
    The traversal coverage is intentionally unchanged, but individual run writes
    are skipped when the run is already Arial 12.
    """
    try:
        style = doc.styles["Normal"]
        if style.font.name != _TARGET_FONT_NAME:
            style.font.name = _TARGET_FONT_NAME

        try:
            current_style_size_pt = (
                round(float(style.font.size.pt), 2)
                if style.font.size is not None
                else None
            )
        except Exception:
            current_style_size_pt = None

        if current_style_size_pt != float(_TARGET_FONT_SIZE_PT):
            style.font.size = _TARGET_FONT_SIZE
    except Exception:
        pass

    for paragraph in doc.paragraphs:
        _normalize_paragraph_runs_font(paragraph)

    for table in doc.tables:
        _normalize_table_runs_font(table)

    for section in doc.sections:
        containers = [
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ]

        for container in containers:
            if container is None:
                continue

            for paragraph in container.paragraphs:
                _normalize_paragraph_runs_font(paragraph)

            for table in container.tables:
                _normalize_table_runs_font(table)


def paragraph_runs_text(paragraph) -> tuple[str, list[tuple[int, int, int]]]:
    """
    Return paragraph full text and positional map.

    The map contains tuples:
    (run_index, start_offset_in_full_text, end_offset_in_full_text)
    """
    pieces: list[str] = []
    positions: list[tuple[int, int, int]] = []

    cursor = 0
    for idx, run in enumerate(paragraph.runs):
        text = run.text or ""
        pieces.append(text)
        start = cursor
        end = start + len(text)
        positions.append((idx, start, end))
        cursor = end

    return "".join(pieces), positions


def find_run_index_at_offset(
    positions: list[tuple[int, int, int]],
    offset: int,
) -> int:
    """
    Find the run index containing the given character offset.
    """
    for run_index, start, end in positions:
        if start <= offset < end:
            return run_index

    if positions and offset == positions[-1][2]:
        return positions[-1][0]

    return -1


def _iter_text_segments_for_range(
    paragraph,
    positions: list[tuple[int, int, int]],
    start: int,
    end: int,
) -> list[tuple[str, object | None]]:
    """
    Return styled text segments for a slice of paragraph full text.

    Each returned item is:
        (text_chunk, source_run_or_none)

    The slice is split on original run boundaries so style can be preserved as
    closely as possible when rebuilding a paragraph after placeholder
    replacements.

    If the paragraph has no runs, a single style-less segment is returned.
    """
    if start >= end:
        return []

    if not paragraph.runs:
        return [(paragraph.text[start:end], None)]

    segments: list[tuple[str, object | None]] = []

    for run_index, run_start, run_end in positions:
        if run_end <= start:
            continue
        if run_start >= end:
            break

        seg_start = max(start, run_start)
        seg_end = min(end, run_end)
        if seg_start >= seg_end:
            continue

        source_run = paragraph.runs[run_index]
        local_start = seg_start - run_start
        local_end = seg_end - run_start
        text_chunk = (source_run.text or "")[local_start:local_end]

        if text_chunk:
            segments.append((text_chunk, source_run))

    return segments


def _append_styled_segment(paragraph, text: str, src_run) -> None:
    """
    Append a run with the given text, copying style from `src_run` when present.
    """
    if not text:
        return

    new_run = paragraph.add_run(text)
    copy_run_style(src_run, new_run)


def _append_range_with_original_styles(
    paragraph,
    original_paragraph,
    positions: list[tuple[int, int, int]],
    start: int,
    end: int,
) -> None:
    """
    Append a text slice from the original paragraph while preserving run-level
    styling across original run boundaries.
    """
    for text_chunk, src_run in _iter_text_segments_for_range(
        original_paragraph,
        positions,
        start,
        end,
    ):
        _append_styled_segment(paragraph, text_chunk, src_run)


def _placeholder_matches_for_paragraph(
    full_text: str,
    mapping: dict[str, str],
) -> list[tuple[int, int, str, str]]:
    """
    Return all placeholder matches for a paragraph.

    Each match is:
        (start_offset, end_offset, placeholder, replacement)

    Only placeholders that exist in `mapping` are returned. Unknown placeholder
    tokens are ignored so they remain unchanged in output.
    """
    if not full_text or _PLACEHOLDER_SENTINEL not in full_text:
        return []

    matches: list[tuple[int, int, str, str]] = []

    for match in _PLACEHOLDER_TOKEN_RE.finditer(full_text):
        placeholder = match.group(0)
        if placeholder not in mapping:
            continue
        matches.append(
            (
                match.start(),
                match.end(),
                placeholder,
                str(mapping[placeholder]),
            )
        )

    return matches


def _replace_placeholders_single_pass_in_paragraph(
    paragraph,
    mapping: dict[str, str],
) -> bool:
    """
    Replace all mapped placeholders in a paragraph in a single rebuild pass.

    This preserves split-run placeholder support while avoiding repeated
    paragraph flattening and repeated nested placeholder loops.

    Styling behavior:
    - text outside placeholders keeps its original run styling
    - replacement text inherits the style of the first participating run of the
      matched placeholder, which matches current behavior
    """
    if not mapping:
        return False

    full_text, positions = paragraph_runs_text(paragraph)
    if not full_text or _PLACEHOLDER_SENTINEL not in full_text:
        return False

    matches = _placeholder_matches_for_paragraph(full_text, mapping)
    if not matches:
        return False

    original_paragraph = paragraph
    original_runs = list(paragraph.runs)

    rebuilt_segments: list[tuple[str, object | None]] = []
    cursor = 0

    for start, end, _placeholder, replacement in matches:
        if cursor < start:
            rebuilt_segments.extend(
                _iter_text_segments_for_range(original_paragraph, positions, cursor, start)
            )

        start_run_idx = find_run_index_at_offset(positions, start)
        replacement_style_run = (
            original_runs[start_run_idx] if start_run_idx >= 0 else None
        )
        rebuilt_segments.append((replacement, replacement_style_run))
        cursor = end

    if cursor < len(full_text):
        rebuilt_segments.extend(
            _iter_text_segments_for_range(
                original_paragraph,
                positions,
                cursor,
                len(full_text),
            )
        )

    clear_paragraph_runs(paragraph)

    for text_chunk, src_run in rebuilt_segments:
        _append_styled_segment(paragraph, text_chunk, src_run)

    return True


def replace_placeholder_once_in_paragraph(
    paragraph,
    placeholder: str,
    replacement: str,
) -> bool:
    """
    Replace one occurrence of a placeholder in a paragraph, even if it is split
    across multiple runs.

    The replacement text inherits the style of the first participating run.
    """
    full_text, positions = paragraph_runs_text(paragraph)
    if not full_text or placeholder not in full_text:
        return False

    start = full_text.find(placeholder)
    end = start + len(placeholder)

    start_run_idx = find_run_index_at_offset(positions, start)
    end_run_idx = find_run_index_at_offset(positions, end - 1)

    if start_run_idx < 0 or end_run_idx < 0:
        return False

    start_run = paragraph.runs[start_run_idx]
    end_run = paragraph.runs[end_run_idx]

    start_run_global_start = positions[start_run_idx][1]
    end_run_global_start = positions[end_run_idx][1]

    prefix = start_run.text[: start - start_run_global_start]
    suffix = end_run.text[(end - end_run_global_start):]

    start_run.text = f"{prefix}{replacement}{suffix}"

    for idx in range(start_run_idx + 1, end_run_idx + 1):
        paragraph.runs[idx].text = ""

    return True


def replace_placeholder_all_in_paragraph(
    paragraph,
    placeholder: str,
    replacement: str,
) -> None:
    """
    Replace all occurrences of a placeholder in a paragraph, robustly across
    runs.

    This function keeps the public contract unchanged. For performance, it uses
    the same single-pass paragraph rewrite engine as the document-level helpers,
    but scoped to a single placeholder mapping.
    """
    _replace_placeholders_single_pass_in_paragraph(paragraph, {placeholder: replacement})


def _replace_placeholders_in_table(table, mapping: dict[str, str]) -> None:
    """
    Replace placeholders in every paragraph of a table.
    """
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _replace_placeholders_single_pass_in_paragraph(paragraph, mapping)


def _replace_placeholders_in_container(container, mapping: dict[str, str]) -> None:
    """
    Replace placeholders in a header/footer-like container.

    Supported content:
    - direct paragraphs
    - table cell paragraphs
    """
    for paragraph in container.paragraphs:
        _replace_placeholders_single_pass_in_paragraph(paragraph, mapping)

    for table in container.tables:
        _replace_placeholders_in_table(table, mapping)


def replace_placeholders_everywhere(doc: Document, mapping: dict[str, str]) -> None:
    """
    Replace placeholders across all document paragraphs and table paragraphs.

    This implementation is optimized to:
    - skip paragraphs without '{{'
    - flatten each candidate paragraph only once
    - replace all placeholders in a paragraph in one rebuild pass
    """
    if not mapping:
        return

    for paragraph in doc.paragraphs:
        _replace_placeholders_single_pass_in_paragraph(paragraph, mapping)

    for table in doc.tables:
        _replace_placeholders_in_table(table, mapping)


def replace_placeholders_in_headers_and_footers(
    doc: Document,
    mapping: dict[str, str],
) -> None:
    """
    Replace placeholders in all header/footer variants for every section:
    - default header/footer
    - first-page header/footer
    - even-page header/footer
    """
    if not mapping:
        return

    for section in doc.sections:
        containers = [
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ]

        for container in containers:
            if container is None:
                continue
            _replace_placeholders_in_container(container, mapping)


def replace_literal_text_everywhere(doc: Document, replacements: dict[str, str]) -> None:
    """
    Replace literal fallback text in document body paragraphs and table cells.

    This is kept for backward compatibility with older templates that still
    contain slash-based wording instead of explicit placeholders.
    """
    for paragraph in doc.paragraphs:
        original = paragraph.text or ""
        updated = original
        for src, dst in replacements.items():
            if src in updated:
                updated = updated.replace(src, dst)
        if updated != original:
            paragraph.text = updated

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    original = paragraph.text or ""
                    updated = original
                    for src, dst in replacements.items():
                        if src in updated:
                            updated = updated.replace(src, dst)
                    if updated != original:
                        paragraph.text = updated


def clear_table_body_keep_header(table, header_rows: int = 1) -> None:
    """
    Remove all table rows after the given header row count.
    """
    while len(table.rows) > header_rows:
        tbl = table._tbl
        tr = table.rows[header_rows]._tr
        tbl.remove(tr)


def set_cell_alignment(
    cell,
    *,
    horizontal: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.CENTER,
    vertical: WD_ALIGN_VERTICAL = WD_ALIGN_VERTICAL.CENTER,
) -> None:
    """
    Apply horizontal and vertical alignment to every paragraph in a cell.
    """
    cell.vertical_alignment = vertical
    for paragraph in cell.paragraphs:
        paragraph.alignment = horizontal


def clone_paragraph_properties(src_paragraph, dst_paragraph) -> None:
    """
    Clone paragraph properties XML including tabs, indents, spacing and
    alignment.
    """
    dst_p = dst_paragraph._p
    src_p = src_paragraph._p

    if dst_p.pPr is not None:
        dst_p.remove(dst_p.pPr)

    if src_p.pPr is not None:
        dst_p.insert(0, deepcopy(src_p.pPr))


def copy_run_style(src_run, dst_run) -> None:
    """
    Copy run-level style from one run to another.
    """
    if src_run is None:
        return

    dst_run.bold = src_run.bold
    dst_run.italic = src_run.italic
    dst_run.underline = src_run.underline
    dst_run.font.name = src_run.font.name
    dst_run.font.size = src_run.font.size

    try:
        if src_run.font.color is not None and src_run.font.color.rgb is not None:
            dst_run.font.color.rgb = src_run.font.color.rgb
    except Exception:
        pass

    try:
        dst_run.font.highlight_color = src_run.font.highlight_color
    except Exception:
        pass


def insert_paragraph_after(paragraph):
    """
    Insert and return a new paragraph immediately after the given paragraph.
    """
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return paragraph._parent.add_paragraph().__class__(new_p, paragraph._parent)


def clear_paragraph_runs(paragraph) -> None:
    """
    Remove all runs from a paragraph while preserving paragraph properties.
    """
    p = paragraph._p
    for child in list(p):
        if child.tag.endswith("}r") or child.tag.endswith("}hyperlink"):
            p.remove(child)


def iter_document_paragraphs(doc: Document) -> Iterable:
    """
    Yield all body paragraphs and all table-cell paragraphs in document order
    blocks.

    This helper is useful for search-style traversal where body + tables are
    both valid search locations.
    """
    for paragraph in doc.paragraphs:
        yield paragraph

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


__all__ = [
    "set_global_font_arial_12",
    "paragraph_runs_text",
    "find_run_index_at_offset",
    "replace_placeholder_once_in_paragraph",
    "replace_placeholder_all_in_paragraph",
    "replace_placeholders_everywhere",
    "replace_placeholders_in_headers_and_footers",
    "replace_literal_text_everywhere",
    "clear_table_body_keep_header",
    "set_cell_alignment",
    "clone_paragraph_properties",
    "copy_run_style",
    "insert_paragraph_after",
    "clear_paragraph_runs",
    "iter_document_paragraphs",
]