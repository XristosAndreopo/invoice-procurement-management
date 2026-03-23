"""
app/reports/award_decision_docx.py

Generate "ΑΠΟΦΑΣΗ ΑΝΑΘΕΣΗΣ" as DOCX bytes using a Word template
and placeholder replacement.

IMPORTANT DOMAIN CHANGE
-----------------------
Handler organizational data must now come from the selected procurement
handler assignment, not by assuming one fixed department/directory on Personnel.

Placeholders supported for the selected handler assignment:
  {{HANDLER_DIRECTORY}}
  {{HANDLER_DEPARTMENT}}

IMPORTANT MASTER-DATA RULE
--------------------------
The placeholder `{{armodiothtas}}` must be resolved from the ALE–KAE master
directory using the Procurement.ale code.

The Procurement entity stores only:
- procurement.ale

It does not store:
- procurement.armodiothtas
- procurement.ale_kae relationship

Therefore this report must look up the ALE master row explicitly through the
shared master-data service.

TEMPLATE ALIGNMENT
------------------
This implementation is aligned to the current final award decision DOCX template.

Confirmed placeholders present in the template include:
- {{SHORT_DATE}}
- {{SERVICE_UNIT_NAME}}
- {{SERVICE_UNIT_PHONE}}
- {{SERVICE_UNIT_REGION}}
- {{COMMANDER_ROLE_TYPE}}
- {{service.commander}}
- {{WINNER_SUPPLIER_LINE}}
- {{HANDLER_DIRECTORY}}
- {{HANDLER_DEPARTMENT}}
- {{ML_TOTAL_WORDS}}
- cost-analysis placeholders
"""

from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Any, Iterable, Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ..services.master_data_service import get_ale_row_by_code
from .common.amounts import money_plain, money_words_el, percent, to_decimal
from .common.docx_utils import (
    clear_table_body_keep_header,
    replace_placeholders_everywhere,
    set_cell_alignment,
    set_global_font_arial_12,
)
from .common.domain import (
    format_recipients_block,
    resolve_document_total,
    resolve_document_total_value,
    resolve_handler_department,
    resolve_handler_directory,
    winner_supplier_line,
)
from .common.formatting import safe_text, short_date_el, upper_no_accents, upper_service_name


def _template_path() -> Path:
    """
    Resolve the DOCX template path for the award decision document.
    """
    app_dir = Path(__file__).resolve().parents[1]
    return app_dir / "templates" / "docx" / "award_decision_template.docx"


def _find_items_table(doc: Document):
    """
    Find the materials/items table in the template.

    Current template contract:
    - 5 columns
    - header includes CPV and ΠΕΡΙΓΡΑΦΗ
    """
    for table in doc.tables:
        if len(table.columns) != 5 or not table.rows:
            continue

        header_text = " ".join(cell.text for cell in table.rows[0].cells).upper()
        if "CPV" in header_text and "ΠΕΡΙΓΡΑΦΗ" in header_text:
            return table

    return None


def _find_cost_table(doc: Document):
    """
    Find the pricing/cost table in the template.

    Current template contract:
    - 6 columns
    - header includes ΤΙΜΗ, ΜΟΝ, ΣΥΝΟΛΟ
    """
    for table in doc.tables:
        if len(table.columns) != 6 or not table.rows:
            continue

        header_text = " ".join(cell.text for cell in table.rows[0].cells).upper()
        if "ΤΙΜΗ" in header_text and "ΜΟΝ" in header_text and "ΣΥΝΟΛΟ" in header_text:
            return table

    return None


def _fill_items_table(table, materials: list[Any]) -> None:
    """
    Populate the item-description table for award decision rendering.
    """
    clear_table_body_keep_header(table, header_rows=1)

    if not materials:
        row = table.add_row().cells
        row[0].text = "—"
        row[1].text = "Δεν υπάρχουν γραμμές υλικών/υπηρεσιών."
        row[2].text = "—"
        row[3].text = "—"
        row[4].text = "—"
        for cell in row:
            set_cell_alignment(
                cell,
                horizontal=WD_ALIGN_PARAGRAPH.CENTER,
                vertical=WD_ALIGN_VERTICAL.CENTER,
            )
        return

    for idx, line in enumerate(materials, start=1):
        row = table.add_row().cells
        row[0].text = str(idx)
        row[1].text = safe_text(getattr(line, "description", None), default="")
        row[2].text = safe_text(getattr(line, "unit", None))
        row[3].text = safe_text(getattr(line, "quantity", None))
        row[4].text = safe_text(getattr(line, "cpv", None))

        for cell in row:
            set_cell_alignment(
                cell,
                horizontal=WD_ALIGN_PARAGRAPH.CENTER,
                vertical=WD_ALIGN_VERTICAL.CENTER,
            )


def _add_cost_summary_row(table, label: str, amount: Any) -> None:
    """
    Append one summary row to the cost-analysis table.
    """
    row = table.add_row()
    merged_cell = row.cells[0].merge(row.cells[4])
    merged_cell.text = label
    row.cells[5].text = money_plain(amount)

    set_cell_alignment(
        merged_cell,
        horizontal=WD_ALIGN_PARAGRAPH.RIGHT,
        vertical=WD_ALIGN_VERTICAL.CENTER,
    )
    set_cell_alignment(
        row.cells[5],
        horizontal=WD_ALIGN_PARAGRAPH.CENTER,
        vertical=WD_ALIGN_VERTICAL.CENTER,
    )


def _fill_cost_table(table, materials: list[Any], analysis: dict[str, Any]) -> None:
    """
    Populate the pricing/cost table and append calculated summary rows.
    """
    clear_table_body_keep_header(table, header_rows=1)

    if not materials:
        row = table.add_row().cells
        for idx in range(6):
            row[idx].text = "—"
        for cell in row:
            set_cell_alignment(
                cell,
                horizontal=WD_ALIGN_PARAGRAPH.CENTER,
                vertical=WD_ALIGN_VERTICAL.CENTER,
            )
    else:
        for idx, line in enumerate(materials, start=1):
            row = table.add_row().cells
            row[0].text = str(idx)
            row[1].text = safe_text(getattr(line, "description", None), default="")
            row[2].text = safe_text(getattr(line, "unit", None))
            row[3].text = safe_text(getattr(line, "quantity", None))
            row[4].text = money_plain(getattr(line, "unit_price", None))
            row[5].text = money_plain(getattr(line, "total_pre_vat", None))

            for cell in row:
                set_cell_alignment(
                    cell,
                    horizontal=WD_ALIGN_PARAGRAPH.CENTER,
                    vertical=WD_ALIGN_VERTICAL.CENTER,
                )

    public_withholdings = analysis.get("public_withholdings") or {}
    income_tax = analysis.get("income_tax") or {}

    public_pct = percent(public_withholdings.get("total_percent", 0))
    income_tax_pct = percent(income_tax.get("rate_percent", 0))
    vat_pct = percent(analysis.get("vat_percent", 0))

    _add_cost_summary_row(table, "Μερικό Σύνολο", analysis.get("sum_total", 0))
    _add_cost_summary_row(
        table,
        f"Κρατήσεις Υπερ Δημοσίου ({public_pct}%)",
        public_withholdings.get("total_amount", 0),
    )
    _add_cost_summary_row(
        table,
        f"ΦΕ ({income_tax_pct}%)",
        income_tax.get("amount", 0),
    )
    _add_cost_summary_row(
        table,
        f"ΦΠΑ ({vat_pct}%)",
        analysis.get("vat_amount", 0),
    )
    _add_cost_summary_row(table, "Τελικό Σύνολο", analysis.get("payable_total", 0))


def _resolve_armodiothtas(procurement: Any) -> str:
    """
    Resolve the responsibility text for the award decision.

    Resolution order:
    1. lookup ALE row by procurement.ale
    2. return ALE responsibility
    3. fallback to "—"
    """
    ale_code = (getattr(procurement, "ale", None) or "").strip()
    if not ale_code:
        return "—"

    ale_row = get_ale_row_by_code(ale_code)
    if ale_row is None:
        return "—"

    responsibility = getattr(ale_row, "responsibility", None)
    return safe_text(responsibility)


def _apply_award_paragraph_vat_text(doc: Document, proc_type: str, vat_percent: Any) -> None:
    """
    Apply the current template's VAT-tail wording rule in the award paragraph.
    """
    vat_is_zero = to_decimal(vat_percent).quantize(to_decimal("0.01")) == to_decimal("0.00")
    replacement_tail = ", άνευ ΦΠΑ." if vat_is_zero else " και ΦΠΑ."

    targets = [
        f", {proc_type}.",
        f", {proc_type} .",
        f"/ {proc_type}.",
        f"/{proc_type}.",
    ]

    for paragraph in doc.paragraphs:
        text = paragraph.text or ""
        if "συμπεριλαμβανομένων κρατήσεων" not in text:
            continue

        updated = text
        for target in targets:
            if target in updated:
                updated = updated.replace(target, replacement_tail)

        updated = updated.replace(", άνευ ΦΠΑ/ και ΦΠΑ.", ", άνευ ΦΠΑ.")
        updated = updated.replace(", άνευ ΦΠΑ / και ΦΠΑ.", ", άνευ ΦΠΑ.")
        updated = updated.replace(", άνευ ΦΠΑ/και ΦΠΑ.", ", άνευ ΦΠΑ.")
        updated = updated.replace(", άνευ ΦΠΑ /και ΦΠΑ.", ", άνευ ΦΠΑ.")

        if updated != text:
            paragraph.text = updated
            break


@dataclass(frozen=True)
class AwardDecisionConstants:
    """
    Future-proof constants container for award decision generation.
    """
    pass


def build_award_decision_docx(
    *,
    procurement: Any,
    service_unit: Any,
    winner: Optional[Any],
    other_suppliers: Iterable[Any],
    analysis: dict,
    is_services: bool,
    constants: Optional[AwardDecisionConstants] = None,
) -> bytes:
    """
    Build the award decision DOCX and return it as bytes.
    """
    _ = constants

    template_path = _template_path()
    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")

    doc = Document(str(template_path))

    proc_type = "παροχή υπηρεσιών" if is_services else "προμήθεια υλικών"

    public_withholdings = analysis.get("public_withholdings") or {}
    income_tax = analysis.get("income_tax") or {}

    public_pct = percent(public_withholdings.get("total_percent", 0))
    income_tax_pct = percent(income_tax.get("rate_percent", 0))
    vat_pct = percent(analysis.get("vat_percent", 0))

    winner_name = safe_text(getattr(winner, "name", None), default="—")
    winner_afm = safe_text(getattr(winner, "afm", None), default="—")
    winner_line = winner_supplier_line(winner)

    mapping: dict[str, str] = {
        "{{SHORT_DATE}}": short_date_el(),
        "{{PROC_TYPE}}": proc_type,
        "{{SERVICE_UNIT_NAME}}": upper_service_name(
            safe_text(getattr(service_unit, "description", None), default="—")
        ),
        "{{SERVICE_UNIT_PHONE}}": safe_text(getattr(service_unit, "phone", None)),
        "{{SERVICE_UNIT_REGION}}": safe_text(getattr(service_unit, "region", None)),
        "{{procurement.aay}}": safe_text(getattr(procurement, "aay", None)),
        "{{procurement.adam_aay}}": safe_text(getattr(procurement, "adam_aay", None)),
        "{{procurement.identity_prosklisis}}": safe_text(
            getattr(procurement, "identity_prosklisis", None)
        ),
        "{{procurement.adam_prosklisis}}": safe_text(getattr(procurement, "adam_prosklisis", None)),
        "{{ procurement.adam_prosklisis}}": safe_text(
            getattr(procurement, "adam_prosklisis", None)
        ),
        "{{procurement.ale}}": safe_text(getattr(procurement, "ale", None)),
        "{{current_year}}": str(getattr(procurement, "fiscal_year", None) or ""),
        "{{current year}}": str(getattr(procurement, "fiscal_year", None) or ""),
        "{{armodiothtas}}": _resolve_armodiothtas(procurement),
        "{{WINNER_SUPPLIER_LINE}}": winner_line,
        "{{supplier.name}}": winner_name,
        "{{supplier.afm}}": winner_afm,
        "{{RECIPIENTS_INFO}}": format_recipients_block(other_suppliers),
        "{{service.commander}}": safe_text(getattr(service_unit, "commander", None), default="—"),
        "{{COMMANDER_ROLE_TYPE}}": safe_text(
            getattr(service_unit, "commander_role_type", None)
        ),
        "{{AN_PUBLIC_WITHHOLD_PERCENT}}": f" ({public_pct}%)",
        "{{AN_PUBLIC_WITHHOLD_TOTAL}}": money_plain(public_withholdings.get("total_amount", 0)),
        "{{AN_INCOME_TAX_RATE}}": f" ({income_tax_pct}%)",
        "{{AN_INCOME_TAX_TOTAL}}": money_plain(income_tax.get("amount", 0)),
        "{{AN_VAT_PERCENT}}": vat_pct,
        "{{AN_VAT_AMOUNT}}": money_plain(analysis.get("vat_amount", 0)),
        "{{AN_SUM_TOTAL}}": money_plain(analysis.get("sum_total", 0)),
        "{{AN_PAYABLE_TOTAL}}": money_plain(analysis.get("payable_total", 0)),
        "{{ML_TOTAL}}": resolve_document_total(procurement, analysis),
        "{{ML_TOTAL_WORDS}}": money_words_el(resolve_document_total_value(procurement, analysis)),
        "{{HANDLER_DIRECTORY}}": upper_no_accents(resolve_handler_directory(procurement)),
        "{{HANDLER_DEPARTMENT}}": resolve_handler_department(procurement),
    }

    replace_placeholders_everywhere(doc, mapping)
    _apply_award_paragraph_vat_text(doc, proc_type, analysis.get("vat_percent", 0))

    materials = list(getattr(procurement, "materials", []) or [])

    items_table = _find_items_table(doc)
    if items_table is not None:
        _fill_items_table(items_table, materials)

    cost_table = _find_cost_table(doc)
    if cost_table is not None:
        _fill_cost_table(cost_table, materials, analysis)

    set_global_font_arial_12(doc)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def build_award_decision_filename(
    *,
    procurement: Any,
    winner: Optional[Any],
    is_services: bool,
) -> str:
    """
    Build a human-readable filename for the generated award decision document.
    """
    kind = "Παροχής Υπηρεσιών" if is_services else "Προμήθειας Υλικών"
    supplier_name = safe_text(getattr(winner, "name", None), default="—")
    total_str = resolve_document_total(procurement, {})

    return f"Απόφαση Ανάθεσης {kind} {supplier_name} {total_str}.docx"